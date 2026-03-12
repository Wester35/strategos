import re
from docx import Document


class DocumentParser:

    def parse(self, file_path: str) -> str:
        doc = Document(file_path)

        paragraphs = [
            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()
        ]

        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    tables.append(" | ".join(cells))

        return "\n".join(paragraphs + tables)