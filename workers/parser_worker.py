import re
from datetime import datetime

from PySide6.QtCore import QThread, Signal
from docx import Document


class ParserWorker(QThread):

    progress_update = Signal(int)
    status_update = Signal(str)
    parsing_finished = Signal(dict)
    error_occurred = Signal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        try:

            self.status_update.emit("Чтение файла...")
            doc = Document(self.file_path)

            paragraphs = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text)

            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))

            full_text = "\n".join(paragraphs + tables_text)

            self.status_update.emit("Парсинг блока 1...")
            block1 = self.parse_block1(full_text)
            self.progress_update.emit(33)

            self.status_update.emit("Парсинг блока 2...")
            block2 = self.parse_block2(full_text)
            self.progress_update.emit(66)

            self.status_update.emit("Парсинг блока 3...")
            block3 = self.parse_block3(full_text)
            self.progress_update.emit(100)

            result = {
                "metadata": {
                    "file_name": self.file_path,
                    "parsed_at": datetime.now().isoformat(),
                    "total_paragraphs": len(paragraphs),
                    "total_tables": len(doc.tables)
                },
                "block1": block1,
                "block2": block2,
                "block3": block3,
                "full_text": full_text
            }

            self.parsing_finished.emit(result)

        except Exception as e:
            self.error_occurred.emit(str(e))

    def parse_block1(self, content):
        block1 = {}

        goal_match = re.search(
            r'\*\*Цель.*?\*\*.*?\n(.*?)(?=\n\n|\n\*\*)',
            content,
            re.DOTALL | re.IGNORECASE
        )

        if goal_match:
            block1["strategic_goal"] = goal_match.group(1).strip()

        targets_match = re.search(
            r'\*\*Конкретные целевые показатели.*?\*\*(.*?)(?=\n\n\*\*|\n*Критерии)',
            content,
            re.DOTALL | re.IGNORECASE
        )

        if targets_match:
            targets_text = targets_match.group(1)

            targets = {}

            for line in targets_text.split('\n'):
                line = line.strip()

                if line and ':' in line and not line.startswith('**'):
                    parts = line.split(':', 1)

                    if len(parts) == 2:
                        key, value = parts
                        targets[key.strip()] = value.strip()

            block1["targets"] = targets

        return block1

    def parse_block2(self, content):
        block2 = {}

        departments = {
            "hr": "Департамент персонала",
            "sales": "Департамент продаж и маркетинга",
            "finance": "Бухгалтерия / Финансовый департамент",
            "design": "Департамент проектирования"
        }

        block2_start = re.search(r'Блок 2\.', content, re.IGNORECASE)

        if block2_start:

            block2_text = content[block2_start.end():]

            for dept_key, dept_name in departments.items():

                dept_data = {}

                dept_pattern = rf'\*\*{re.escape(dept_name)}.*?\*\*'

                dept_match = re.search(dept_pattern, block2_text, re.IGNORECASE)

                if dept_match:

                    start = dept_match.end()
                    end = len(block2_text)

                    for next_name in departments.values():

                        if next_name != dept_name:

                            next_pattern = rf'\*\*{re.escape(next_name)}.*?\*\*'

                            next_match = re.search(
                                next_pattern,
                                block2_text[start:],
                                re.IGNORECASE
                            )

                            if next_match:
                                end = start + next_match.start()
                                break

                    dept_section = block2_text[start:end]

                    role_match = re.search(
                        r'\*\*Главная роль:\*\*\s*(.*?)(?=\n\n|\n\*\*)',
                        dept_section,
                        re.DOTALL
                    )

                    if role_match:
                        dept_data["main_role"] = role_match.group(1).strip()

                    goals_match = re.search(
                        r'\*\*Цели на 12 месяцев:\*\*\s*(.*?)(?=\n\n\*\*|\nМетрики)',
                        dept_section,
                        re.DOTALL
                    )

                    if goals_match:

                        goals_text = goals_match.group(1)

                        goals = []

                        for line in goals_text.split('\n'):
                            line = line.strip()

                            if line and re.match(r'^\d+\.', line):
                                goals.append(re.sub(r'^\d+\.\s*', '', line))

                        dept_data["goals"] = goals

                    block2[dept_key] = dept_data

        return block2

    def parse_block3(self, content):

        block3 = {}

        departments = {
            "hr": "Департамент 1: Персонал",
            "sales": "Департамент 2: Продажи и маркетинг",
            "finance": "Департамент 3: Бухгалтерия",
            "design": "Департамент 4: Проектирование"
        }

        block3_start = re.search(r'Блок 3\.', content, re.IGNORECASE)

        if block3_start:

            block3_text = content[block3_start.end():]

            for dept_key, dept_name in departments.items():

                dept_data = {}

                dept_pattern = rf'\*\*{re.escape(dept_name)}.*?\*\*'

                dept_match = re.search(dept_pattern, block3_text, re.IGNORECASE)

                if dept_match:

                    start = dept_match.end()
                    end = len(block3_text)

                    for next_name in departments.values():

                        if next_name != dept_name:

                            next_pattern = rf'\*\*{re.escape(next_name)}.*?\*\*'

                            next_match = re.search(
                                next_pattern,
                                block3_text[start:],
                                re.IGNORECASE
                            )

                            if next_match:
                                end = start + next_match.start()
                                break

                    dept_section = block3_text[start:end]

                    for q_num in range(1, 5):

                        quarter_key = f"Q{q_num}"
                        quarter_data = {}

                        quarter_pattern = rf'\*\*Q{q_num} .*?\*\*'

                        quarter_match = re.search(
                            quarter_pattern,
                            dept_section,
                            re.IGNORECASE
                        )

                        if quarter_match:

                            quarter_data["found"] = True
                            dept_data[quarter_key] = quarter_data

                    block3[dept_key] = dept_data

        return block3